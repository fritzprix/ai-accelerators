import pypandoc
import os
import sys
import argparse
from docx import Document

def create_reference_template():
    """Pandoc에서 사용할 기본 참조 템플릿 제작 (스타일 커스터마이징용)"""
    template_path = 'template_reference.docx'
    if not os.path.exists(template_path):
        doc = Document()
        doc.add_heading('Heading 1', 0)
        doc.add_heading('Heading 2', 1)
        doc.add_paragraph('Normal Text with standard styling.')
        doc.save(template_path)
    return template_path

def convert_md_to_docx(input_md):
    output_docx = input_md.replace('.md', '.docx')
    template = 'template_reference.docx'
    
    if not os.path.exists(template):
        create_reference_template()
        
    print(f"Converting {input_md} to {output_docx}...")
    
    # Standalone mode, with Table of Contents and reference template
    try:
        pypandoc.convert_file(
            input_md, 
            'docx', 
            outputfile=output_docx,
            extra_args=[
                '--standalone', 
                '--toc', 
                f'--reference-doc={template}',
                '--syntax-highlighting'
            ]
        )
        print(f"✅ Success: {output_docx}")
    except Exception as e:
        print(f"❌ Error: {e}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("file", help="Markdown file to convert")
    args = parser.parse_args()
    
    if os.path.exists(args.file):
        convert_md_to_docx(args.file)
    else:
        print(f"File not found: {args.file}")
